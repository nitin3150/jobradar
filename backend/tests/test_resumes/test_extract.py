from pathlib import Path

from app.resumes.extract import extract_text, MAX_EXTRACT_CHARS


def test_extract_txt(tmp_path: Path):
    f = tmp_path / "r.txt"
    f.write_text("Nitin — ML Engineer, LangGraph, FastAPI")
    assert "LangGraph" in extract_text(f, "text/plain")


def test_extract_docx(tmp_path: Path):
    import docx

    f = tmp_path / "r.docx"
    doc = docx.Document()
    doc.add_paragraph("Senior ML Engineer")
    doc.add_paragraph("Skills: PyTorch, FastAPI")
    doc.save(str(f))
    out = extract_text(
        f,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "PyTorch" in out and "Senior ML Engineer" in out


def test_unsupported_type_returns_empty(tmp_path: Path):
    f = tmp_path / "r.doc"
    f.write_bytes(b"\xff\xfe legacy doc")
    assert extract_text(f, "application/msword") == ""


def test_corrupt_pdf_returns_empty(tmp_path: Path):
    f = tmp_path / "bad.pdf"
    f.write_bytes(b"not really a pdf")
    # Failure must never raise — extraction is best-effort.
    assert extract_text(f, "application/pdf") == ""


def test_output_is_capped(tmp_path: Path):
    f = tmp_path / "big.txt"
    f.write_text("x" * (MAX_EXTRACT_CHARS + 5000))
    assert len(extract_text(f, "text/plain")) == MAX_EXTRACT_CHARS
